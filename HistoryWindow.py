import sqlite3
import customtkinter as ctk
from users_db import SqliteDB
import pandas as pd
from reportlab.lib.pagesizes import letter
from reportlab.pdfgen import canvas
import os
from datetime import datetime
import subprocess
from tkinter import filedialog
import calendar
import json


class ConsumptionHistoryWindow:
    def __init__(self, parent, width, height, abonent_id, title="Учет коммунальных услуг АО_Корммаш",
                 resizable=(False, False), icon='image/korm.ico'):
        self.root = ctk.CTkToplevel(parent)
        self.root.title(title)
        self.root.geometry(f"{width}x{height}")
        self.root.resizable(resizable[0], resizable[1])
        if icon:
            self.root.iconbitmap(icon)

        self.abonent_id = abonent_id
        self.last_pdf_path = None  # Будем хранить путь к последнему созданному PDF
        print(f"Тип abonent_id: {type(self.abonent_id)}, значение: {self.abonent_id}")

        # Создаем основной контейнер
        self.main_frame = ctk.CTkScrollableFrame(self.root)
        self.main_frame.pack(fill=ctk.BOTH, expand=True, padx=20, pady=20)

        # Заголовок
        title_frame = ctk.CTkFrame(self.main_frame)
        title_frame.pack(fill=ctk.X, pady=(0, 20))
        ctk.CTkLabel(title_frame, 
                    text="История потребления",
                    font=("Roboto", 20, "bold")).pack(pady=10)

        # Фрейм для выбора периода
        period_frame = ctk.CTkFrame(self.main_frame)
        period_frame.pack(fill=ctk.X, pady=10)

        # Месяц
        month_frame = ctk.CTkFrame(period_frame)
        month_frame.pack(side=ctk.LEFT, padx=10, expand=True)
        
        ctk.CTkLabel(month_frame, 
                    text="📅 Месяц:",
                    font=("Roboto", 14)).pack(pady=5)

        self.months = ["Январь", "Февраль", "Март", "Апрель", "Май", "Июнь",
                      "Июль", "Август", "Сентябрь", "Октябрь", "Ноябрь", "Декабрь"]
        self.month_var = ctk.StringVar(value=self.months[datetime.now().month - 1])
        self.month_combobox = ctk.CTkComboBox(month_frame, 
                                            values=self.months,
                                            variable=self.month_var,
                                            font=("Roboto", 12),
                                            height=35)
        self.month_combobox.pack(pady=5)

        # Год
        year_frame = ctk.CTkFrame(period_frame)
        year_frame.pack(side=ctk.LEFT, padx=10, expand=True)
        
        ctk.CTkLabel(year_frame, 
                    text="📅 Год:",
                    font=("Roboto", 14)).pack(pady=5)
        
        self.year_entry = ctk.CTkEntry(year_frame,
                                     width=100,
                                     height=35,
                                     font=("Roboto", 12))
        self.year_entry.pack(pady=5)
        self.year_entry.insert(0, str(datetime.now().year))

        # Кнопки управления
        buttons_frame = ctk.CTkFrame(self.main_frame)
        buttons_frame.pack(fill=ctk.X, pady=10)

        self.load_button = ctk.CTkButton(buttons_frame, 
                                       text="📊 Загрузить данные",
                                       font=("Roboto", 12),
                                       height=40,
                                       command=self.load_data)
        self.load_button.pack(side=ctk.LEFT, padx=10, pady=5)

        self.calc_button = ctk.CTkButton(buttons_frame, 
                                       text="🧮 Рассчитать потребление",
                                       font=("Roboto", 12),
                                       height=40,
                                       state='disabled',
                                       command=self.calculate_consumption)
        self.calc_button.pack(side=ctk.LEFT, padx=10, pady=5)

        self.generate_registry_button = ctk.CTkButton(buttons_frame, 
                                                    text="📝 Создать реестр",
                                                    font=("Roboto", 12),
                                                    height=40,
                                                    state='disabled',
                                                    command=self.generate_registry)
        self.generate_registry_button.pack(side=ctk.LEFT, padx=10, pady=5)

        self.open_word_button = ctk.CTkButton(buttons_frame, 
                                            text="📄 Открыть Word",
                                            font=("Roboto", 12),
                                            height=40,
                                            state='disabled',
                                            fg_color='green',
                                            command=self.open_word)
        self.open_word_button.pack(side=ctk.LEFT, padx=10, pady=5)

        # Таблица данных
        table_frame = ctk.CTkFrame(self.main_frame)
        table_frame.pack(fill=ctk.BOTH, expand=True, pady=10)
        
        ctk.CTkLabel(table_frame, 
                    text="Данные за период",
                    font=("Roboto", 14)).pack(pady=5)
        
        self.table = ctk.CTkTextbox(table_frame, 
                                  width=550,
                                  height=200,
                                  font=("Roboto", 12))
        self.table.pack(pady=5, padx=10, fill=ctk.BOTH, expand=True)

        # Результаты расчетов
        calc_frame = ctk.CTkFrame(self.main_frame)
        calc_frame.pack(fill=ctk.BOTH, expand=True, pady=10)
        
        ctk.CTkLabel(calc_frame, 
                    text="Результаты расчетов",
                    font=("Roboto", 14)).pack(pady=5)
        
        self.calculation_result = ctk.CTkTextbox(calc_frame, 
                                               width=550,
                                               height=200,
                                               font=("Roboto", 12))
        self.calculation_result.pack(pady=5, padx=10, fill=ctk.BOTH, expand=True)

        self.root.grab_set()
        self.root.focus_set()

        self.transformation_ratio = 1  # Коэффициент трансформации по умолчанию

    def open_word(self):
        """Открывает созданный Word файл в стандартном просмотрщике"""
        if self.last_doc_path and os.path.exists(self.last_doc_path):
            try:
                if os.name == 'nt':  # Для Windows
                    os.startfile(self.last_doc_path)
                elif os.name == 'posix':  # Для Mac и Linux
                    subprocess.run(
                        ['open', self.last_doc_path] if sys.platform == 'darwin' else ['xdg-open', self.last_doc_path])
                self.calculation_result.insert("end", f"\nWord файл открыт: {self.last_doc_path}\n")
            except Exception as e:
                self.calculation_result.insert("end", f"\nОшибка при открытии Word файла: {str(e)}\n")
        else:
            self.calculation_result.insert("end", "\nWord файл не найден. Сначала создайте реестр.\n")

    def get_month_number(self, month_name):
        """Возвращает номер месяца по его названию"""
        return self.months.index(month_name) + 1

    def load_data(self):
        """Загружает данные о потреблении за указанный расчетный месяц и отображает их в интерфейсе"""
        try:
            # 1. Получаем и валидируем параметры периода
            month_name = self.month_var.get()
            month = self.get_month_number(month_name)
            year = int(self.year_entry.get())

            # Валидация введенных значений
            if year < 2000 or year > datetime.now().year + 1:
                raise ValueError("Год должен быть в диапазоне 2000-" + str(datetime.now().year + 1))

            print(f"Параметры запроса: abonent_id={self.abonent_id}, месяц: {month}/{year}")

            # 2. Подключаемся к базе данных
            db = SqliteDB()

            try:
                # Проверка существования таблицы abonents
                if not db.execute_query("SELECT name FROM sqlite_master WHERE type='table' AND name='abonents'",
                                        fetch_mode='one'):
                    raise Exception("Таблица abonents не существует")

                # Получаем имя абонента и коэффициент трансформации
                abonent_data = db.execute_query(
                    "SELECT fulname, transformation_ratio_value FROM abonents WHERE id = ?",
                    (self.abonent_id,),
                    fetch_mode='one'
                )

                if not abonent_data:
                    raise Exception(f"Абонент с ID {self.abonent_id} не найден")

                fulname = abonent_data[0]
                self.transformation_ratio = abonent_data[1] if abonent_data[1] is not None else 1
                print(f"Найден абонент: {fulname}, коэффициент трансформации: {self.transformation_ratio}")

                # 3. Проверяем существование таблицы monthly_data
                table_check = "SELECT name FROM sqlite_master WHERE type='table' AND name='monthly_data'"
                if not db.execute_query(table_check, fetch_mode='one'):
                    self.table.delete("1.0", "end")
                    self.table.insert("end", "Ошибка: таблица monthly_data не существует\n")
                    return


                # 4. Получаем данные за расчетный месяц и предыдущий месяц
                # Для расчета потребления нам нужны показания на начало и конец периода

                # Текущий расчетный месяц (конец периода)
                end_month_data = db.execute_query(
                    "SELECT * FROM monthly_data WHERE abonent_id = ? AND month = ? AND year = ?",
                    (self.abonent_id, month, year),
                    fetch_mode='one'
                )

                # Предыдущий месяц (начало периода)
                prev_month = month - 1 if month > 1 else 12
                prev_year = year if month > 1 else year - 1

                start_month_data = db.execute_query(
                    "SELECT * FROM monthly_data WHERE abonent_id = ? AND month = ? AND year = ?",
                    (self.abonent_id, prev_month, prev_year),
                    fetch_mode='one'
                )

                self.table.delete("1.0", "end")

                if not end_month_data:
                    self.table.insert("end", f"Нет данных за расчетный месяц {month}/{year}\n")
                    return

                if not start_month_data:
                    self.table.insert("end", f"Нет данных за предыдущий месяц {prev_month}/{prev_year}\n")
                    self.table.insert("end", "Расчет будет выполнен от нуля\n")

                # Формируем заголовки таблицы
                headers = ["Параметр", "Предыдущий месяц", "Текущий месяц", "Потребление"]
                self.table.insert("end", "\t".join(headers) + "\n")
                self.table.insert("end", "-" * 70 + "\n")

                # Электроэнергия
                if len(end_month_data) > 4 and end_month_data[4] is not None:
                    prev_value = float(start_month_data[4]) if start_month_data and len(start_month_data) > 4 and \
                                                               start_month_data[4] is not None else 0.0
                    curr_value = float(end_month_data[4])
                    consumption = curr_value - prev_value

                    row = [
                        "Электроэнергия (кВт·ч)",
                        f"{prev_value:.2f}" if start_month_data else "нет данных",
                        f"{curr_value:.2f}",
                        f"{consumption:.2f}"
                    ]
                    self.table.insert("end", "\t".join(row) + "\n")

                # Вода
                if len(end_month_data) > 5 and end_month_data[5] is not None:
                    prev_value = float(start_month_data[5]) if start_month_data and len(start_month_data) > 5 and \
                                                               start_month_data[5] is not None else 0.0
                    curr_value = float(end_month_data[5])
                    consumption = curr_value - prev_value

                    row = [
                        "Вода (м³)",
                        f"{prev_value:.2f}" if start_month_data else "нет данных",
                        f"{curr_value:.2f}",
                        f"{consumption:.2f}"
                    ]
                    self.table.insert("end", "\t".join(row) + "\n")

                # Сточные воды
                if len(end_month_data) > 6 and end_month_data[6] is not None:
                    prev_value = float(start_month_data[6]) if start_month_data and len(start_month_data) > 6 and \
                                                               start_month_data[6] is not None else 0.0
                    curr_value = float(end_month_data[6])
                    consumption = curr_value - prev_value

                    row = [
                        "Сточные воды (м³)",
                        f"{prev_value:.2f}" if start_month_data else "нет данных",
                        f"{curr_value:.2f}",
                        f"{consumption:.2f}"
                    ]
                    self.table.insert("end", "\t".join(row) + "\n")

                # Газ
                if len(end_month_data) > 7 and end_month_data[7] is not None:
                    prev_value = float(start_month_data[7]) if start_month_data and len(start_month_data) > 7 and \
                                                               start_month_data[7] is not None else 0.0
                    curr_value = float(end_month_data[7])
                    consumption = curr_value - prev_value

                    row = [
                        "Газ (м³)",
                        f"{prev_value:.2f}" if start_month_data else "нет данных",
                        f"{curr_value:.2f}",
                        f"{consumption:.2f}"
                    ]
                    self.table.insert("end", "\t".join(row) + "\n")

                # Активируем кнопки после успешной загрузки данных
                self.calc_button.configure(state="normal", fg_color="#1f6aa5")
                self.generate_registry_button.configure(state="normal", fg_color="#1f6aa5")

            finally:
                # Всегда закрываем соединение с БД
                db.close_connection()

        except ValueError as ve:
            self.table.delete("1.0", "end")
            self.table.insert("end", f"Ошибка ввода: {str(ve)}\n")
        except sqlite3.Error as dbe:
            self.table.delete("1.0", "end")
            self.table.insert("end", f"Ошибка базы данных: {str(dbe)}\n")
        except Exception as e:
            self.table.delete("1.0", "end")
            self.table.insert("end", f"Неожиданная ошибка: {str(e)}\n")
            import traceback
            traceback.print_exc()

    def calculate_consumption(self):
        """Вычисляет общее потребление ресурсов за расчетный месяц"""
        try:
            # Проверяем, что данные загружены
            if not self.table.get("1.0", "end-1c"):
                self.calculation_result.delete("1.0", "end")
                self.calculation_result.insert("end", "❌ Ошибка: сначала загрузите данные!\n")
                return

            month_name = self.month_var.get()
            month = self.get_month_number(month_name)
            year = int(self.year_entry.get())

            db = SqliteDB()
            try:
                # Получаем имя абонента
                fulname = db.execute_query(
                    "SELECT fulname FROM abonents WHERE id = ?",
                    (self.abonent_id,),
                    fetch_mode='one'
                )[0]

                # Получаем данные за расчетный месяц и предыдущий месяц
                end_month_data = db.execute_query(
                    "SELECT * FROM monthly_data WHERE abonent_id = ? AND month = ? AND year = ?",
                    (self.abonent_id, month, year),
                    fetch_mode='one'
                )

                # Предыдущий месяц (начало периода)
                prev_month = month - 1 if month > 1 else 12
                prev_year = year if month > 1 else year - 1

                start_month_data = db.execute_query(
                    "SELECT * FROM monthly_data WHERE abonent_id = ? AND month = ? AND year = ?",
                    (self.abonent_id, prev_month, prev_year),
                    fetch_mode='one'
                )

                if not end_month_data:
                    self.calculation_result.delete("1.0", "end")
                    self.calculation_result.insert("end", f"Нет данных за расчетный месяц {month}/{year}\n")
                    return

                # Рассчитываем потребление по каждой услуге
                result_text = f"Расчет потребления за {month_name} {year} года\n"
                result_text += f"Абонент: {fulname}\n\n"

                # Электроэнергия
                if len(end_month_data) > 4 and end_month_data[4] is not None:
                    prev_value = float(start_month_data[4]) if start_month_data and len(start_month_data) > 4 and \
                                                               start_month_data[4] is not None else 0.0
                    curr_value = float(end_month_data[4])
                    consumption = (curr_value - prev_value) * self.transformation_ratio

                    if self.transformation_ratio != 1:
                        result_text += f"Электроэнергия: {consumption:.2f} кВт·ч (с учетом Кт={self.transformation_ratio})\n"
                    else:
                        result_text += f"Электроэнергия: {consumption:.2f} кВт·ч\n"

                # Вода
                if len(end_month_data) > 5 and end_month_data[5] is not None:
                    prev_value = float(start_month_data[5]) if start_month_data and len(start_month_data) > 5 and \
                                                               start_month_data[5] is not None else 0.0
                    curr_value = float(end_month_data[5])
                    consumption = curr_value - prev_value
                    result_text += f"Вода: {consumption:.2f} м³\n"

                # Сточные воды
                if len(end_month_data) > 6 and end_month_data[6] is not None:
                    prev_value = float(start_month_data[6]) if start_month_data and len(start_month_data) > 6 and \
                                                               start_month_data[6] is not None else 0.0
                    curr_value = float(end_month_data[6])
                    consumption = curr_value - prev_value
                    result_text += f"Сточные воды: {consumption:.2f} м³\n"

                # Газ
                if len(end_month_data) > 7 and end_month_data[7] is not None:
                    prev_value = float(start_month_data[7]) if start_month_data and len(start_month_data) > 7 and \
                                                               start_month_data[7] is not None else 0.0
                    curr_value = float(end_month_data[7])
                    consumption = curr_value - prev_value
                    result_text += f"Газ: {consumption:.2f} м³\n"

                # Выводим результат
                self.calculation_result.delete("1.0", "end")
                self.calculation_result.insert("end", result_text)

            finally:
                db.close_connection()

        except Exception as e:
            self.calculation_result.delete("1.0", "end")
            self.calculation_result.insert("end", f"Ошибка расчета: {str(e)}\n")

    def generate_registry(self):
        """Генерирует реестр в формате Word с показаниями и потреблением"""
        try:
            # Проверяем, что данные загружены и рассчитаны
            if not self.table.get("1.0", "end-1c") or not self.calculation_result.get("1.0", "end-1c"):
                self.calculation_result.delete("1.0", "end")
                self.calculation_result.insert("end", "❌ Ошибка: сначала загрузите и рассчитайте данные!\n")
                return

            month_name = self.month_var.get()
            month = self.get_month_number(month_name)
            year = int(self.year_entry.get())

            db = SqliteDB()
            try:
                # Получаем данные абонента с проверкой на None
                abonent_data = db.execute_query(
                    "SELECT fulname FROM abonents WHERE id = ?",
                    (self.abonent_id,),
                    fetch_mode='one'
                )

                if not abonent_data:
                    self.calculation_result.delete("1.0", "end")
                    self.calculation_result.insert("end", "❌ Ошибка: абонент не найден!\n")
                    return

                fulname = abonent_data[0] if abonent_data[0] else "Не указано"

                # Получаем данные за расчетный месяц с проверкой
                end_month_data = db.execute_query(
                    "SELECT * FROM monthly_data WHERE abonent_id = ? AND month = ? AND year = ?",
                    (self.abonent_id, month, year),
                    fetch_mode='one'
                )

                if not end_month_data:
                    self.calculation_result.delete("1.0", "end")
                    self.calculation_result.insert("end", f"❌ Нет данных за {month_name} {year} года!\n")
                    return

                # Получаем данные за предыдущий месяц
                prev_month = month - 1 if month > 1 else 12
                prev_year = year if month > 1 else year - 1

                start_month_data = db.execute_query(
                    "SELECT * FROM monthly_data WHERE abonent_id = ? AND month = ? AND year = ?",
                    (self.abonent_id, prev_month, prev_year),
                    fetch_mode='one'
                )

                # Создаем документ Word
                from docx import Document
                from docx.shared import Pt, Inches
                from docx.enum.text import WD_ALIGN_PARAGRAPH

                doc = Document()

                # Настройка стилей
                style = doc.styles['Normal']
                style.font.name = 'Times New Roman'
                style.font.size = Pt(12)

                # Заголовок
                title = doc.add_paragraph()
                title.alignment = WD_ALIGN_PARAGRAPH.CENTER
                title_run = title.add_run('РЕЕСТР\nвозмещения затрат за потребление электроэнергии и воды\n')
                title_run.bold = True
                title_run.font.size = Pt(14)

                # Период
                period = doc.add_paragraph()
                period.alignment = WD_ALIGN_PARAGRAPH.CENTER
                period.add_run(f'за {month_name} {year} г.\n').bold = True

                # Организация
                org = doc.add_paragraph()
                org.alignment = WD_ALIGN_PARAGRAPH.CENTER
                org.add_run(f'{fulname}\n').bold = True

                # Электроэнергия (индекс 4 в monthly_data)
                if len(end_month_data) > 4 and end_month_data[4] is not None:
                    prev_value = float(start_month_data[4]) if start_month_data and len(start_month_data) > 4 and \
                                                               start_month_data[4] is not None else 0.0
                    curr_value = float(end_month_data[4])
                    consumption = curr_value - prev_value
                    total_consumption = consumption * (
                        self.transformation_ratio if hasattr(self, 'transformation_ratio') else 1)

                    doc.add_paragraph('1. Показания счетчика электроэнергии:', style='Normal')
                    doc.add_paragraph(f'   - на начало периода: {prev_value:.1f} кВт·ч', style='Normal')
                    doc.add_paragraph(f'   - на конец периода: {curr_value:.1f} кВт·ч', style='Normal')
                    doc.add_paragraph(f'   - итого потребление: {consumption:.1f} кВт·ч', style='Normal')

                    if hasattr(self, 'transformation_ratio') and self.transformation_ratio != 1:
                        doc.add_paragraph(f'   - коэффициент трансформации: {self.transformation_ratio}',
                                          style='Normal')
                        doc.add_paragraph(f'   - итого потребление с учетом КТ: {total_consumption:.1f} кВт·ч',
                                          style='Normal')

                    doc.add_paragraph('2. Тариф за потребленную электроэнергию: ______________ руб./кВт·ч',
                                      style='Normal')
                    doc.add_paragraph('   ИТОГО к оплате: ______________________ руб.', style='Normal')
                    doc.add_paragraph('3. Тариф за заявленную мощность: _______________ руб./кВт', style='Normal')
                    doc.add_paragraph('   Заявленная мощность: _________________ кВт', style='Normal')
                    doc.add_paragraph('   ИТОГО к оплате: ________________ руб.', style='Normal')
                    doc.add_paragraph('   ВСЕГО к оплате (п.2 + п.3): ________________ руб.', style='Normal')
                    doc.add_paragraph()

                # Вода (индекс 5)
                if len(end_month_data) > 5 and end_month_data[5] is not None:
                    prev_value = float(start_month_data[5]) if start_month_data and len(start_month_data) > 5 and \
                                                               start_month_data[5] is not None else 0.0
                    curr_value = float(end_month_data[5])
                    consumption = curr_value - prev_value

                    doc.add_paragraph('4. Показания счетчика воды:', style='Normal')
                    doc.add_paragraph(f'   - на начало периода: {prev_value:.1f} м³', style='Normal')
                    doc.add_paragraph(f'   - на конец периода: {curr_value:.1f} м³', style='Normal')
                    doc.add_paragraph(f'   - итого потребление: {consumption:.1f} м³', style='Normal')
                    doc.add_paragraph('   Тариф: ______________ руб./м³', style='Normal')
                    doc.add_paragraph('   ИТОГО к оплате: ______________________ руб.', style='Normal')
                    doc.add_paragraph()

                # Сточные воды (индекс 6)
                if len(end_month_data) > 6 and end_month_data[6] is not None:
                    prev_value = float(start_month_data[6]) if start_month_data and len(start_month_data) > 6 and \
                                                               start_month_data[6] is not None else 0.0
                    curr_value = float(end_month_data[6])
                    consumption = curr_value - prev_value

                    doc.add_paragraph('5. Водоотведение:', style='Normal')
                    doc.add_paragraph(f'{consumption:.1f} м³', style='Normal')
                    doc.add_paragraph('   Тариф: ______________ руб./м³', style='Normal')
                    doc.add_paragraph('   ИТОГО к оплате: ______________________ руб.', style='Normal')
                    doc.add_paragraph()

                # Газ (индекс 7)
                if len(end_month_data) > 7 and end_month_data[7] is not None:
                    prev_value = float(start_month_data[7]) if start_month_data and len(start_month_data) > 7 and \
                                                               start_month_data[7] is not None else 0.0
                    curr_value = float(end_month_data[7])
                    consumption = curr_value - prev_value

                    doc.add_paragraph('6. Показания счетчика газа:', style='Normal')
                    doc.add_paragraph(f'   - на начало периода: {prev_value:.1f} м³', style='Normal')
                    doc.add_paragraph(f'   - на конец периода: {curr_value:.1f} м³', style='Normal')
                    doc.add_paragraph(f'   - итого потребление: {consumption:.1f} м³', style='Normal')
                    doc.add_paragraph('   Тариф: ______________ руб./м³', style='Normal')
                    doc.add_paragraph('   ИТОГО к оплате: ______________________ руб.', style='Normal')
                    doc.add_paragraph()

                # Подписи
                # Загружаем настройки
                settings_file = os.path.join(os.path.dirname(os.path.abspath(__file__)), 'data', 'settings.json')
                try:
                    if os.path.exists(settings_file):
                        with open(settings_file, 'r', encoding='utf-8') as f:
                            settings = json.load(f)
                    else:
                        settings = {}
                except Exception as e:
                    print(f"Ошибка при загрузке настроек: {e}")
                    settings = {}
                
                # Добавляем подписи из настроек
                signatures = settings.get("signatures", [])
                for signature in signatures:
                    position = signature.get("position", "")
                    name = signature.get("name", "")
                    if position and name:
                        doc.add_paragraph(f'{position} {name}\t/____________/', style='Normal')
                
                doc.add_paragraph('Согласовано:', style='Normal')
                doc.add_paragraph('Арендатор ___________________/________________/', style='Normal')

                # Сохраняем документ
                # Получаем название месяца в именительном падеже
                month_names = {
                    1: "январь", 2: "февраль", 3: "март", 4: "апрель",
                    5: "май", 6: "июнь", 7: "июль", 8: "август",
                    9: "сентябрь", 10: "октябрь", 11: "ноябрь", 12: "декабрь"
                }
                month_folder = month_names[month]
                
                # Получаем путь сохранения из настроек
                save_path = settings.get("save_path", r"C:\Реестры по абонентам")
                
                # Создаем подпапку с названием месяца
                month_folder_path = os.path.join(save_path, month_folder)
                os.makedirs(month_folder_path, exist_ok=True)

                # Создаем имя файла без запрещенных символов
                import re
                safe_name = re.sub(r'[\\/*?:"<>|]', "", fulname)
                file_name = f"{safe_name}_{month_name}_{year}_реестр.docx"
                file_path = os.path.join(month_folder_path, file_name)
                doc.save(file_path)

                self.calculation_result.delete("1.0", "end")
                self.calculation_result.insert("end", f"✅ Реестр успешно создан:\n{file_path}\n")

                # Сохраняем путь для открытия
                self.last_doc_path = file_path
                self.open_word_button.configure(state="normal", fg_color="green", text="Открыть Word")


            except Exception as e:
                self.calculation_result.delete("1.0", "end")
                self.calculation_result.insert("end", f"❌ Ошибка создания реестра: {str(e)}\n")
                import traceback
                traceback.print_exc()

            finally:
                db.close_connection()

        except Exception as e:
            self.calculation_result.delete("1.0", "end")
            self.calculation_result.insert("end", f"❌ Ошибка: {str(e)}\n")
